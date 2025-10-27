import google.generativeai as genai
import os
from dotenv import load_dotenv
import json
import base64
import io
from docx import Document

load_dotenv()

# Configure Gemini
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
genai.configure(api_key=GEMINI_API_KEY)


def extract_resume_data(text_message, media_content=None, media_type=None):
    """
    Extract resume data using Gemini AI
    
    Args:
        text_message: Text content from WhatsApp
        media_content: Binary content of media file (PDF/image)
        media_type: MIME type of media
    
    Returns:
        dict: Extracted resume data
    """
    
    extraction_prompt = """
    You are a resume data extraction AI. Extract the following information from the resume or text provided.
    Return ONLY a valid JSON object with these fields:
    
    {
        "name": "full name of the person",
        "email": "email address",
        "phone": "phone/contact number",
        "linkedin": "LinkedIn profile URL if available",
        "skills": ["skill1", "skill2", "skill3"],
        "experience": "years of experience or key role/position",
        "education": "highest degree or institution",
        "summary": "brief 1-line summary of the candidate"
    }
    
    Rules:
    - Return ONLY the JSON object, no markdown, no explanations
    - If a field is not found, use "N/A" for strings or [] for arrays
    - Extract skills as an array
    - Be flexible with resume formats (PDF, image, text)
    - For phone numbers, preserve the format as given
    """
    
    try:
        # Use Gemini 2.0 Flash for images/PDFs and text
        if media_content and media_type:
            print(f"📄 Processing {media_type} file...")
            
            # Use gemini-2.0-flash-exp (latest model with vision capabilities)
            model = genai.GenerativeModel('gemini-2.0-flash-exp')
            
            # Prepare the content based on media type
            if 'image' in media_type:
                # Handle image
                import PIL.Image
                image = PIL.Image.open(io.BytesIO(media_content))
                response = model.generate_content([extraction_prompt, image])
            
            elif 'pdf' in media_type:
                # Handle PDF - convert to base64 and send
                base64_pdf = base64.b64encode(media_content).decode('utf-8')
                response = model.generate_content([
                    extraction_prompt,
                    {
                        "mime_type": "application/pdf",
                        "data": base64_pdf
                    }
                ])
            
            elif 'wordprocessingml' in media_type or 'msword' in media_type or 'docx' in media_type:
                # Handle .docx/.doc files - extract text first
                print("📝 Processing Word document...")
                try:
                    doc_file = io.BytesIO(media_content)
                    doc = Document(doc_file)
                    
                    # Extract all text from document
                    docx_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            docx_text.append(paragraph.text)
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    docx_text.append(cell.text)
                    
                    full_text = '\n'.join(docx_text)
                    print(f"📄 Extracted {len(full_text)} characters from Word document")
                    
                    # Now process with Gemini
                    response = model.generate_content(f"{extraction_prompt}\n\nResume Content:\n{full_text}")
                
                except Exception as e:
                    print(f"❌ Error reading .docx: {e}")
                    # Fallback to text message if available
                    if text_message:
                        response = model.generate_content(f"{extraction_prompt}\n\nResume Content:\n{text_message}")
                    else:
                        return None
            
            else:
                # Try as text for other formats
                model = genai.GenerativeModel('gemini-2.0-flash-exp')
                response = model.generate_content(f"{extraction_prompt}\n\nResume Content:\n{text_message}")
        
        else:
            # Text only
            print("📝 Processing text message...")
            model = genai.GenerativeModel('gemini-2.0-flash-exp')
            response = model.generate_content(f"{extraction_prompt}\n\nResume Content:\n{text_message}")
        
        # Parse response
        result_text = response.text.strip()
        
        # Remove markdown code blocks if present
        if result_text.startswith('```'):
            result_text = result_text.split('```')[1]
            if result_text.startswith('json'):
                result_text = result_text[4:]
        
        # Parse JSON
        extracted_data = json.loads(result_text)
        
        print(f"✅ Successfully extracted data")
        return extracted_data
        
    except json.JSONDecodeError as e:
        print(f"❌ JSON parsing error: {e}")
        print(f"Response was: {result_text}")
        
        # Return a basic structure with error info
        return {
            "name": "N/A",
            "email": "N/A",
            "phone": "N/A",
            "linkedin": "N/A",
            "skills": [],
            "experience": "N/A",
            "education": "N/A",
            "summary": f"Extraction failed: {str(e)}"
        }
        
    except Exception as e:
        print(f"❌ Error extracting resume data: {e}")
        return None


def extract_from_text_only(text):
    """
    Simple extraction from plain text (backup method)
    """
    try:
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        
        prompt = f"""
        Extract contact details from this text and return as JSON:
        {text}
        
        Return format:
        {{"name": "...", "email": "...", "phone": "...", "skills": [], "summary": "..."}}
        """
        
        response = model.generate_content(prompt)
        result_text = response.text.strip()
        
        # Clean markdown
        if '```' in result_text:
            result_text = result_text.split('```')[1]
            if result_text.startswith('json'):
                result_text = result_text[4:]
        
        return json.loads(result_text)
        
    except Exception as e:
        print(f"Error in text extraction: {e}")
        return None


if __name__ == "__main__":
    # Test the extraction
    test_text = """
    John Doe
    Software Engineer
    Email: john.doe@example.com
    Phone: +1-555-0123
    LinkedIn: linkedin.com/in/johndoe
    
    Skills: Python, JavaScript, React, Node.js, Machine Learning
    Experience: 5 years in software development
    Education: B.S. Computer Science, MIT
    """
    
    result = extract_resume_data(test_text)
    print("\nTest Extraction Result:")
    print(json.dumps(result, indent=2))

